"""AI Proposal Generation router.

POST   /api/v1/ai/proposals/generate           — AI-generate & persist a proposal
GET    /api/v1/ai/proposals                    — list tenant proposals (?lead_id, ?status)
GET    /api/v1/ai/proposals/{id}               — get one proposal
PATCH  /api/v1/ai/proposals/{id}               — update title / status / sections
DELETE /api/v1/ai/proposals/{id}               — delete
GET    /api/v1/ai/proposals/{id}/export/docx   — download as Word (.docx)
GET    /api/v1/ai/proposals/{id}/export/html   — download styled HTML (Ctrl+P → PDF)
"""

from __future__ import annotations

import io
import logging
from typing import List, Optional
from uuid import UUID

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import HTMLResponse, StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from core.dependencies import get_current_user, get_db
from models.lead import Lead
from models.proposal import Proposal
from models.user import User
from schemas.proposal import (
    ProposalGenerateRequest,
    ProposalResponse,
    ProposalUpdateRequest,
)
from services import proposal_service

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ai", tags=["AI"])


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


async def _enrich(p: Proposal, db: AsyncSession) -> ProposalResponse:
    """Attach lead display fields to a ProposalResponse."""
    lead = await db.get(Lead, p.lead_id)
    return ProposalResponse(
        id=p.id,
        tenant_id=p.tenant_id,
        lead_id=p.lead_id,
        created_by=p.created_by,
        proposal_type=p.proposal_type,
        title=p.title,
        status=p.status,
        ai_tone=p.ai_tone,
        content_markdown=p.content_markdown,
        content_html=p.content_html,
        sections=p.sections,
        created_at=p.created_at,
        updated_at=p.updated_at,
        lead_company=lead.company_name if lead else None,
        lead_contact=lead.contact_name if lead else None,
    )


async def _get_or_404(
    proposal_id: UUID, db: AsyncSession, tenant_id: UUID
) -> Proposal:
    p = await db.get(Proposal, proposal_id)
    if not p or p.tenant_id != tenant_id:
        raise HTTPException(status_code=404, detail="Proposal not found")
    return p


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------


@router.post("/proposals/generate", response_model=ProposalResponse, status_code=201)
async def generate_proposal(
    req: ProposalGenerateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Generate an AI-powered proposal for a lead and save it as a draft."""
    lead = await db.get(Lead, req.lead_id)
    if not lead or lead.tenant_id != current_user.tenant_id:
        raise HTTPException(status_code=404, detail="Lead not found")

    sections = await proposal_service.generate_proposal(
        db=db,
        tenant_id=current_user.tenant_id,
        user_id=current_user.id,
        lead_id=req.lead_id,
        proposal_type=req.proposal_type,
        tone=req.tone,
        product_description=req.product_description,
    )

    title = sections.get("title") or f"Proposal for {lead.company_name}"
    content_md = proposal_service.sections_to_markdown(sections, lead.company_name or "")
    content_html = proposal_service.sections_to_html(sections, lead.company_name or "")

    proposal = Proposal(
        tenant_id=current_user.tenant_id,
        lead_id=req.lead_id,
        created_by=current_user.id,
        proposal_type=req.proposal_type,
        title=title,
        status="draft",
        ai_tone=req.tone,
        content_markdown=content_md,
        content_html=content_html,
        sections=sections,
    )
    db.add(proposal)
    await db.commit()
    await db.refresh(proposal)
    logger.info("[proposal] created id=%s lead=%s", proposal.id, lead.id)
    return await _enrich(proposal, db)


@router.get("/proposals", response_model=List[ProposalResponse])
async def list_proposals(
    lead_id: Optional[UUID] = Query(None),
    status: Optional[str] = Query(None),
    limit: int = Query(50, le=200),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List proposals for the current tenant, optionally filtered by lead or status."""
    q = (
        select(Proposal)
        .where(Proposal.tenant_id == current_user.tenant_id)
        .order_by(Proposal.created_at.desc())
        .limit(limit)
        .offset(offset)
    )
    if lead_id:
        q = q.where(Proposal.lead_id == lead_id)
    if status:
        q = q.where(Proposal.status == status)
    rows = (await db.execute(q)).scalars().all()
    return [await _enrich(p, db) for p in rows]


@router.get("/proposals/{proposal_id}", response_model=ProposalResponse)
async def get_proposal(
    proposal_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    p = await _get_or_404(proposal_id, db, current_user.tenant_id)
    return await _enrich(p, db)


@router.patch("/proposals/{proposal_id}", response_model=ProposalResponse)
async def update_proposal(
    proposal_id: UUID,
    req: ProposalUpdateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update title, status, or sections of a proposal."""
    p = await _get_or_404(proposal_id, db, current_user.tenant_id)

    if req.title is not None:
        p.title = req.title
    if req.status is not None:
        valid_statuses = {"draft", "sent", "accepted", "rejected"}
        if req.status not in valid_statuses:
            raise HTTPException(400, f"Invalid status. Choose from: {valid_statuses}")
        p.status = req.status
    if req.content_markdown is not None:
        p.content_markdown = req.content_markdown
    if req.sections is not None:
        p.sections = req.sections
        # Re-render HTML from the updated sections so export stays in sync
        lead = await db.get(Lead, p.lead_id)
        company = lead.company_name if lead else ""
        p.content_html = proposal_service.sections_to_html(req.sections, company)

    await db.commit()
    await db.refresh(p)
    return await _enrich(p, db)


@router.delete("/proposals/{proposal_id}", status_code=204)
async def delete_proposal(
    proposal_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    p = await _get_or_404(proposal_id, db, current_user.tenant_id)
    await db.delete(p)
    await db.commit()


@router.get("/proposals/{proposal_id}/export/docx")
async def export_proposal_docx(
    proposal_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Export proposal as a Word (.docx) file — download immediately."""
    try:
        from docx import Document  # type: ignore[import]
        from docx.shared import Pt, RGBColor  # type: ignore[import]
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="python-docx is not installed. Word export is unavailable.",
        )

    p = await _get_or_404(proposal_id, db, current_user.tenant_id)
    lead = await db.get(Lead, p.lead_id)
    company = lead.company_name if lead else "Client"
    sections = p.sections or {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title block ───────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    run = title_para.add_run(sections.get("title", p.title))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    title_para.paragraph_format.space_after = Pt(6)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Prepared for: {company}   |   Prepared by: AveonApex")
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    meta.paragraph_format.space_after = Pt(14)
    doc.add_paragraph("─" * 80)

    # ── Section helper ────────────────────────────────────────────────────
    def _add_section(heading: str, content: str) -> None:
        h = doc.add_heading(heading, level=2)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        h.paragraph_format.space_before = Pt(18)
        for block in str(content or "").split("\n\n"):
            block = block.strip()
            if block:
                p_obj = doc.add_paragraph(block)
                p_obj.paragraph_format.space_after = Pt(6)

    def _add_list_section(heading: str, items: list, style: str = "List Bullet") -> None:
        h = doc.add_heading(heading, level=2)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        h.paragraph_format.space_before = Pt(18)
        for item in items:
            doc.add_paragraph(str(item), style=style)

    # ── Sections ──────────────────────────────────────────────────────────
    _add_section("Executive Summary", sections.get("executive_summary", ""))
    _add_section("Scope of Work", sections.get("scope_of_work", ""))
    _add_list_section("Key Deliverables", sections.get("key_deliverables", []))

    # Pricing table
    pricing = sections.get("pricing", [])
    if pricing:
        h = doc.add_heading("Pricing", level=2)
        if h.runs:
            h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
        h.paragraph_format.space_before = Pt(18)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, label in enumerate(["Service", "Description", "Investment"]):
            hdr[i].text = label
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for row in pricing:
            cells = tbl.add_row().cells
            cells[0].text = str(row.get("item", ""))
            cells[1].text = str(row.get("description", ""))
            cells[2].text = str(row.get("price", ""))

    _add_section("Timeline", sections.get("timeline", ""))
    _add_section("Why AveonApex", sections.get("why_us", ""))

    # Next steps as numbered list
    next_raw = sections.get("next_steps", "")
    if next_raw:
        parts: list[str] = []
        for sep in ("•", "\n-", "\n•"):
            next_raw = next_raw.replace(sep, "\n")
        parts = [
            p_text.strip().lstrip("-•1234567890.)").strip()
            for p_text in next_raw.split("\n")
            if p_text.strip().lstrip("-•1234567890.)").strip()
        ]
        if not parts:
            parts = [s.strip() for s in next_raw.split(".") if s.strip()]
        _add_list_section("Next Steps", parts, style="List Number")

    _add_section("Terms & Conditions", sections.get("terms", ""))

    # Footer
    doc.add_paragraph("─" * 80).paragraph_format.space_before = Pt(24)
    footer = doc.add_paragraph(
        "This proposal is confidential and valid for 30 days from the date of issue. © AveonApex"
    )
    if footer.runs:
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = "".join(
        c if c.isalnum() or c in " _-" else "_" for c in p.title[:50]
    ).strip()
    filename = f"proposal_{safe_name}.docx"

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/proposals/{proposal_id}/export/html", response_class=HTMLResponse)
async def export_proposal_html(
    proposal_id: UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Return the styled HTML proposal. Open in a browser and use Ctrl+P to save as PDF."""
    p = await _get_or_404(proposal_id, db, current_user.tenant_id)
    if p.content_html:
        return HTMLResponse(content=p.content_html)
    # Regenerate if the HTML column is empty (older rows before this feature)
    lead = await db.get(Lead, p.lead_id)
    company = lead.company_name if lead else ""
    html = proposal_service.sections_to_html(p.sections or {}, company)
    return HTMLResponse(content=html)
